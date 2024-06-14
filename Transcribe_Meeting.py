# transcribe_meeting.py
import openai
from docx import Document

# Set up your OpenAI API key
openai.api_key = 'your api key' # give your api key 

# Define the transcription function
def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']

# Define the meeting minutes function
def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

# Placeholder functions for extraction and analysis
def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. ..."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. ..."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. ..."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. ..."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)

# Main workflow
if __name__ == "__main__":
    audio_file_path = "/home/ubuntu/Downloads/ya.wav"
    transcription = transcribe_audio(audio_file_path)
    minutes = meeting_minutes(transcription)
    print(minutes)

    save_as_docx(minutes, 'meeting_minutes.docx')

